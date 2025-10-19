#!/usr/bin/env python3
"""
Fill a DOCX template with resume content from a Markdown file while preserving template styles.
Usage:
  python fill_template_resume.py --template <template.docx> --md <resume.md> --out <output.docx>
"""
import argparse
from pathlib import Path
from typing import Optional
from docx import Document  # factory function to load a .docx
from docx.document import Document as DocxDocument  # type for annotations
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def clear_body(document: DocxDocument):
    body = document._element.body
    for child in list(body):
        body.remove(child)


def add_paragraph(document: DocxDocument, text: str, style: Optional[str] = None, align_center: bool = False):
    p = document.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            # Fallback if style not present in template
            pass
    run = p.add_run(text)
    if style is None:
        run.font.size = Pt(11)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullet(document: DocxDocument, text: str):
    p = document.add_paragraph(text)
    try:
        p.style = 'List Bullet'
    except Exception:
        pass
    for run in p.runs:
        run.font.size = Pt(11)


def fill_from_markdown(document: DocxDocument, md_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()

    for i, raw in enumerate(lines):
        line = raw.rstrip()
        if not line:
            continue
        # Headings
        if line.startswith('# '):
            add_paragraph(document, line[2:].strip(), style='Heading 1', align_center=True)
        elif line.startswith('## '):
            add_paragraph(document, line[3:].strip(), style='Heading 2')
        elif line.startswith('### '):
            add_paragraph(document, line[4:].strip(), style='Heading 3')
        # Horizontal rules or --- separators
        elif line.strip() == '---':
            # Add a small spacer
            document.add_paragraph('')
        # Bullets
        elif line.startswith(('- ', '* ', '• ')):
            add_bullet(document, line[2:].strip())
        # Markdown bold-only line: **Text** -> bold paragraph
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            p = document.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        else:
            # Strip markdown bold decoration in-flow
            text = line.replace('**', '')
            add_paragraph(document, text)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--template', required=True)
    ap.add_argument('--md', required=True)
    ap.add_argument('--out', required=True)
    args = ap.parse_args()

    template_path = Path(args.template)
    md_path = Path(args.md)
    out_path = Path(args.out)

    if not template_path.exists():
        raise SystemExit(f"Template not found: {template_path}")
    if not md_path.exists():
        raise SystemExit(f"Markdown file not found: {md_path}")

    doc: DocxDocument = Document(str(template_path))

    # Clear the body content but keep styles, headers, and footers
    clear_body(doc)

    # Fill with markdown content
    fill_from_markdown(doc, md_path)

    # Save
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved template-filled resume to: {out_path}")


if __name__ == '__main__':
    main()
