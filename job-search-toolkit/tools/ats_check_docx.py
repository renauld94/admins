#!/usr/bin/env python3
"""
ATS Checker for DOCX resumes.
- Scans for ATS blockers (tables, images, headers/footers-heavy content)
- Checks structural sections
- Validates presence of contact info
- Measures keyword coverage vs. job description
- Outputs a Markdown report
"""

import re
import sys
import json
from pathlib import Path
from collections import Counter
from datetime import datetime

from docx import Document

STOPWORDS = set(
    """
a an and are as at be but by for from has have in into is it its of on or that the to was were will with your you their our this those these they them themself oneself itself itselfs over under within without across about above below behind before after again further then once here there when where why how all any both each few more most other some such no nor not only own same so than too very can just don should now 
    """.split()
)

SECTION_HINTS = [
    "summary", "professional summary", "profile", "executive profile",
    "experience", "work experience", "professional experience",
    "education", "skills", "technical skills", "achievements", "key achievements",
    "projects", "languages", "certifications"
]

KEY_PHRASES_HINTS = [
    "ai", "artificial intelligence", "ml", "machine learning",
    "analytics", "data", "director", "leadership", "governance",
    "cloud", "aws", "databricks", "kubernetes", "spark", "airflow",
    "predictive", "pipeline", "feature", "reliability", "uptime",
    "compliance", "hipaa", "team", "stakeholder"
]

PHONE_RE = re.compile(r"(\+?\d[\d\s\-()]{7,}\d)")
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # Include text in tables as they might hide content from ATS
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    # Headers/footers (presence only; extracting text is limited)
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                parts.append(p.text)
        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                parts.append(p.text)
    return "\n".join([t for t in parts if t])


def tokenize(text: str):
    text = text.lower()
    text = re.sub(r"[^a-z0-9+]+", " ", text)
    tokens = [t for t in text.split() if t and t not in STOPWORDS and len(t) > 2]
    return tokens


def extract_keywords_from_job(job_text: str, top_n: int = 30):
    tokens = tokenize(job_text)
    # keep alphanum terms, plus curated hints
    counts = Counter(tokens)
    top = [w for w, _ in counts.most_common(top_n)]
    # merge with hints
    merged = list(dict.fromkeys(top + KEY_PHRASES_HINTS))
    return merged


def analyze_structure(doc: Document):
    issues = []
    tables = len(doc.tables)
    images = getattr(doc.inline_shapes, 'shape_id', None)
    num_images = len(doc.inline_shapes) if hasattr(doc, 'inline_shapes') else 0

    sections_found = set()
    headings = [p.text.strip().lower() for p in doc.paragraphs if p.style and 'Heading' in p.style.name]
    for h in headings:
        for hint in SECTION_HINTS:
            if hint in h:
                sections_found.add(hint)

    # Simple heuristic for lists
    bullets = sum(1 for p in doc.paragraphs if (p.style and 'List' in p.style.name) or p.text.strip().startswith(tuple(['-', '•', '*'])))

    # Headers/footers presence
    headers = sum(1 for s in doc.sections if s.header and s.header.paragraphs)
    footers = sum(1 for s in doc.sections if s.footer and s.footer.paragraphs)

    if tables > 0:
        issues.append(f"Uses {tables} table(s) — ATS may struggle parsing table content.")
    if num_images > 0:
        issues.append(f"Contains {num_images} image(s) — ATS ignores images; ensure no key info is image-only.")
    if headers > 0 or footers > 0:
        issues.append("Has header/footer content — keep minimal; some ATS skip these.")
    if bullets == 0:
        issues.append("No bullet lists detected — ATS and recruiters prefer scannable bullets.")

    return {
        "tables": tables,
        "images": num_images,
        "headers": headers,
        "footers": footers,
        "bullets": bullets,
        "sections_found": sorted(list(sections_found)),
        "issues": issues,
    }


def keyword_coverage(resume_text: str, keywords: list[str]):
    text = resume_text.lower()
    present = []
    missing = []
    for kw in keywords:
        if kw in text:
            present.append(kw)
        else:
            missing.append(kw)
    return present, missing


def detect_contact_info(text: str):
    phones = PHONE_RE.findall(text)
    emails = EMAIL_RE.findall(text)
    return list(set(phones)), list(set(emails))


def score_report(struct, present_keywords, total_keywords):
    # Structure base: start at 60, deduct for issues, add for sections
    score = 60
    # penalize blockers
    if struct['tables'] > 0:
        score -= min(10, struct['tables'] * 3)
    if struct['images'] > 0:
        score -= min(5, struct['images'])
    if struct['headers'] > 0 or struct['footers'] > 0:
        score -= 3
    if struct['bullets'] == 0:
        score -= 5
    # reward sections
    reward_sections = len(struct['sections_found'])
    score += min(10, reward_sections)

    # Keyword coverage: add up to 30
    cov = (len(present_keywords) / max(1, total_keywords))
    score += int(30 * cov)

    return max(0, min(100, score)), int(100 * cov)


def main():
    if len(sys.argv) < 3:
        print("Usage: ats_check_docx.py <resume.docx> <job_description.txt> [output_md]")
        sys.exit(1)

    resume_path = Path(sys.argv[1])
    job_path = Path(sys.argv[2])
    output_md = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    if not resume_path.exists():
        print(f"❌ Resume not found: {resume_path}")
        sys.exit(2)
    if not job_path.exists():
        print(f"❌ Job description not found: {job_path}")
        sys.exit(2)

    doc = Document(str(resume_path))
    resume_text = read_docx_text(resume_path)
    job_text = job_path.read_text(encoding='utf-8', errors='ignore')

    # Structure
    struct = analyze_structure(doc)

    # Keywords
    keywords = extract_keywords_from_job(job_text, top_n=35)
    present, missing = keyword_coverage(resume_text, keywords)

    # Contact
    phones, emails = detect_contact_info(resume_text)

    # Score
    overall, cov_pct = score_report(struct, present, len(keywords))

    # Prepare report
    ts = datetime.now().strftime('%Y-%m-%d %H:%M')
    lines = []
    lines.append(f"# ATS Check Report\n")
    lines.append(f"Generated: {ts}\n")
    lines.append("")
    lines.append(f"**Resume:** {resume_path.name}")
    lines.append(f"**Job Description:** {job_path.name}")
    lines.append("")
    lines.append(f"## Overall Scores")
    lines.append(f"- Overall ATS Readiness: **{overall}/100**")
    lines.append(f"- Keyword Coverage: **{cov_pct}%** ({len(present)}/{len(keywords)})")
    lines.append("")
    lines.append("## Structure Analysis")
    lines.append(f"- Tables: {struct['tables']}")
    lines.append(f"- Images: {struct['images']}")
    lines.append(f"- Bullets detected: {struct['bullets']}")
    lines.append(f"- Headers: {struct['headers']}, Footers: {struct['footers']}")
    lines.append(f"- Sections found: {', '.join(struct['sections_found']) if struct['sections_found'] else 'None'}")
    if struct['issues']:
        lines.append("- Issues:")
        for iss in struct['issues']:
            lines.append(f"  - {iss}")
    else:
        lines.append("- Issues: None detected")
    lines.append("")

    # Contact
    lines.append("## Contact Info Detected")
    lines.append(f"- Phones: {', '.join(phones) if phones else 'None found'}")
    lines.append(f"- Emails: {', '.join(emails) if emails else 'None found'}")
    lines.append("")

    # Keywords
    lines.append("## Keyword Coverage")
    lines.append("- Present:")
    for kw in sorted(present):
        lines.append(f"  - {kw}")
    lines.append("")
    lines.append("- Missing (consider weaving in if relevant):")
    for kw in sorted(missing):
        lines.append(f"  - {kw}")
    lines.append("")

    report_text = "\n".join(lines)

    if not output_md:
        output_md = resume_path.parent / (resume_path.stem + "_ATS_REPORT.md")
    Path(output_md).write_text(report_text, encoding='utf-8')
    print(f"\n✓ ATS report saved to: {output_md}")


if __name__ == "__main__":
    main()
