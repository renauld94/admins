#!/usr/bin/env python3
"""
Exact Resume Recreator - Build resume from scratch matching original formatting
Uses python-docx with precise measurements to recreate 2-page layout
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import sys

def create_exact_resume(output_path):
    """Create resume with exact formatting matching original template"""
    
    doc = Document()
    
    # ==================== PAGE SETUP ====================
    # Match original: Letter size, tight margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    
    # ==================== HEADER SECTION ====================
    # Name - Large, Bold, Centered
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run("SIMON RENAULD")
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name.paragraph_format.space_after = Pt(2)
    
    # Contact Info - Smaller, Centered
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("Ho Chi Minh City, Vietnam | +84 964 294 885 | sn.renauld@gmail.com")
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(11)
    contact.paragraph_format.space_after = Pt(6)
    
    # LinkedIn (optional)
    linkedin = doc.add_paragraph()
    linkedin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linkedin_run = linkedin.add_run("linkedin.com/in/simon-renauld")
    linkedin_run.font.name = 'Calibri'
    linkedin_run.font.size = Pt(10)
    linkedin.paragraph_format.space_after = Pt(10)
    
    # ==================== HELPER FUNCTIONS ====================
    
    def add_section_header(text):
        """Add blue section header with underline"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 120)  # Dark blue
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Add underline by creating a border
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1F4E78"/></w:pBdr>' % nsdecls('w'))
        pPr.append(pBdr)
        return p
    
    def add_paragraph_tight(text, bold=False, indent=0, size=11):
        """Add paragraph with tight spacing"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)  # Tight line spacing
        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        if bold:
            run.font.bold = True
        return p
    
    def add_bullet(text, size=10):
        """Add bullet point with tight spacing"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(12)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(f"• {text}")
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        return p
    
    def add_job_header(title, company, dates):
        """Add job title, company, dates in specific format"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)
        
        # Title - Bold
        title_run = p.add_run(title)
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        
        # Separator
        p.add_run(" | ")
        
        # Company - Bold
        company_run = p.add_run(company)
        company_run.font.name = 'Calibri'
        company_run.font.size = Pt(11)
        company_run.font.bold = True
        
        # Dates - Italic, right-aligned would be complex, so we'll add inline
        p.add_run(" | ")
        dates_run = p.add_run(dates)
        dates_run.font.name = 'Calibri'
        dates_run.font.size = Pt(10)
        dates_run.font.italic = True
        
        return p
    
    # ==================== CONTENT ====================
    
    # PROFESSIONAL SUMMARY
    add_section_header("Professional Summary")
    add_paragraph_tight(
        "Strategic analytics leader with 15+ years driving business transformation through data innovation. "
        "Proven track record building high-performing teams, implementing scalable AI/ML solutions, and delivering "
        "measurable ROI in healthcare, government, and corporate sectors. Expert in cloud architecture (AWS, GCP, Azure), "
        "advanced analytics, and cross-functional leadership across APAC and North American markets."
    )
    
    # CORE COMPETENCIES
    add_section_header("Core Competencies")
    competencies = [
        "AI/ML Strategy & Implementation",
        "Data Architecture & Engineering", 
        "Cloud Solutions (AWS, Azure, GCP)",
        "Team Leadership & Development",
        "Healthcare Analytics & BI",
        "Predictive Modeling & Forecasting",
        "ETL/ELT Pipeline Development",
        "Stakeholder Engagement",
        "Agile Project Management",
        "Data Governance & Security"
    ]
    # Create 2-column layout using table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for i, comp in enumerate(competencies):
        row = i % 5
        col = i // 5
        cell = table.rows[row].cells[col]
        cell.text = f"• {comp}"
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
    
    # Remove table borders
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # PROFESSIONAL EXPERIENCE
    add_section_header("Professional Experience")
    
    # Job 1: Jio Health
    add_job_header("Director of Data Analytics & Engineering", "Jio Health", "2022 – Present")
    add_bullet("Lead 15-member analytics team delivering enterprise BI, predictive modeling, and cloud data solutions across Vietnam")
    add_bullet("Architected AWS-based data lake processing 2M+ patient records, reducing query times by 60%")
    add_bullet("Implemented ML models for patient churn prediction (85% accuracy) and treatment outcome forecasting")
    add_bullet("Established data governance framework ensuring HIPAA/GDPR compliance across all systems")
    add_bullet("Developed executive dashboards driving $2M+ in operational cost savings through resource optimization")
    
    # Job 2: Freelance
    add_job_header("Senior Data Consultant", "Freelance / Contract", "2020 – 2022")
    add_bullet("Delivered 12+ analytics projects for healthcare, fintech, and government clients across APAC region")
    add_bullet("Built real-time streaming pipelines using Kafka, Spark, and Snowflake for financial transaction monitoring")
    add_bullet("Designed PowerBI dashboards for Canadian government agencies, improving decision-making efficiency by 40%")
    add_bullet("Mentored junior data engineers and analysts, establishing best practices for code quality and documentation")
    
    # Job 3: Laval University
    add_job_header("Research Scientist & GIS Analyst", "Laval University", "2015 – 2020")
    add_bullet("Conducted spatial analysis and machine learning research on environmental datasets (10TB+)")
    add_bullet("Published 8 peer-reviewed papers on geospatial modeling and climate prediction algorithms")
    add_bullet("Developed Python-based geospatial tools adopted by 200+ researchers internationally")
    add_bullet("Secured $500K+ in research funding through competitive grant applications")
    
    # EDUCATION
    add_section_header("Education")
    add_paragraph_tight("Master of Science in Geography (Geomatics & Remote Sensing)", bold=True)
    add_paragraph_tight("Laval University, Quebec, Canada | 2013 – 2015", size=10)
    add_paragraph_tight("")  # Small gap
    add_paragraph_tight("Bachelor of Science in Geography", bold=True)
    add_paragraph_tight("Laval University, Quebec, Canada | 2010 – 2013", size=10)
    
    # KEY ACHIEVEMENTS
    add_section_header("Key Achievements")
    add_bullet("Reduced data processing costs by 45% through cloud infrastructure optimization (AWS, Databricks)")
    add_bullet("Increased team productivity 30% by implementing CI/CD pipelines and automated testing frameworks")
    add_bullet("Awarded 'Innovation Leader of the Year' at Jio Health for AI-driven predictive analytics platform")
    
    # LANGUAGES
    add_section_header("Languages")
    add_paragraph_tight("English (Fluent) | Vietnamese (Basic) | French (Native)", size=10)
    
    # TECHNICAL SKILLS
    add_section_header("Technical Skills")
    add_paragraph_tight(
        "Python, R, SQL, Scala | TensorFlow, PyTorch, Scikit-learn | AWS (S3, EC2, Lambda, Redshift), "
        "Azure (Data Factory, Synapse), GCP (BigQuery) | Databricks, Snowflake, Apache Spark | "
        "PowerBI, Tableau, Looker | PostgreSQL, MySQL, MongoDB | Docker, Kubernetes, Git, Airflow",
        size=10
    )
    
    # ==================== SAVE ====================
    print(f"Creating resume: {output_path}")
    doc.save(output_path)
    print("✓ DOCX created with precise formatting")
    
    return output_path

def convert_to_pdf(docx_path):
    """Convert to PDF using LibreOffice"""
    import subprocess
    from pathlib import Path
    
    output_dir = Path(docx_path).parent
    pdf_path = str(docx_path).replace('.docx', '.pdf')
    
    print("Converting to PDF...")
    result = subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', str(output_dir), str(docx_path)
    ], capture_output=True, text=True, timeout=30)
    
    if Path(pdf_path).exists():
        print(f"✓ PDF created: {pdf_path}")
        
        # Check page count
        info_result = subprocess.run(
            ['pdfinfo', pdf_path],
            capture_output=True, text=True
        )
        for line in info_result.stdout.split('\n'):
            if 'Pages:' in line:
                print(f"  {line.strip()}")
        
        return pdf_path
    else:
        print("✗ PDF conversion failed")
        return None

if __name__ == "__main__":
    output_path = sys.argv[1] if len(sys.argv) > 1 else "Simon_Renauld_Resume_RECREATED.docx"
    
    print("="*70)
    print("EXACT RESUME RECREATION - FROM SCRATCH")
    print("="*70)
    
    docx_path = create_exact_resume(output_path)
    pdf_path = convert_to_pdf(docx_path)
    
    print("\n" + "="*70)
    print("✓ RECREATION COMPLETE")
    print("="*70)
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path}")
    print("="*70)
    print("\nReview the PDF and compare with original!")
