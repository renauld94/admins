#!/usr/bin/env python3
"""
DOCX Template Filler v2 - Direct Content Replacement
Preserves exact template formatting while replacing text content
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def read_content_file(content_path):
    """Read and parse the content guide file"""
    with open(content_path, 'r', encoding='utf-8') as f:
        return f.read()

def clear_document_content(doc):
    """Clear all text content from document while preserving structure"""
    # Clear all paragraphs
    for paragraph in doc.paragraphs:
        paragraph.clear()
    
    # Clear all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.clear()

def fill_template_smart(template_path, content_text, output_path):
    """Fill DOCX template with new content, preserving formatting"""
    
    print(f"Loading template: {template_path}")
    doc = Document(template_path)
    
    # Parse content into sections
    lines = content_text.strip().split('\n')
    content_lines = [line for line in lines if line.strip()]
    
    print(f"Content lines to add: {len(content_lines)}")
    
    # Clear existing content
    print("Clearing template content...")
    # Clear all paragraph text but keep structure
    para_count = len(doc.paragraphs)
    for i in range(para_count - 1, -1, -1):
        p = doc.paragraphs[i]
        # Don't delete, just clear
        for run in p.runs:
            run.text = ""
    
    # Now add content to first paragraphs, create new ones as needed
    print("Adding new content...")
    
    para_index = 0
    for line in content_lines:
        line = line.strip()
        if not line:
            continue
            
        # Get or create paragraph
        if para_index < len(doc.paragraphs):
            p = doc.paragraphs[para_index]
        else:
            p = doc.add_paragraph()
        
        # Determine formatting based on content
        if line.isupper() and len(line) > 10 and not line.startswith('•'):
            # Major section header
            p.style = 'Heading 1'
            p.add_run(line).bold = True
        elif line.endswith(':') or (line.isupper() and len(line) < 50):
            # Subsection header
            p.style = 'Heading 2'
            p.add_run(line).bold = True
        elif line.startswith('•'):
            # Bullet point
            p.style = 'List Bullet'
            p.add_run(line)
        else:
            # Normal text
            p.style = 'Normal'
            # Check if should be bold (job titles, company names, dates)
            if re.match(r'^[A-Z][^.]*\d{4}', line) or re.match(r'^[A-Z][A-Za-z\s&,]+$', line):
                p.add_run(line).bold = True
            else:
                p.add_run(line)
        
        para_index += 1
    
    # Save output
    print(f"Saving to: {output_path}")
    doc.save(output_path)
    print("✓ DOCX created successfully")
    
    return output_path

def convert_to_pdf(docx_path):
    """Convert DOCX to PDF using LibreOffice"""
    import subprocess
    
    output_dir = Path(docx_path).parent
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    print("Converting to PDF...")
    cmd = [
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', str(output_dir),
        str(docx_path)
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
    
    if Path(pdf_path).exists():
        print(f"✓ PDF created: {pdf_path}")
        return pdf_path
    else:
        print(f"✗ PDF conversion failed")
        print(result.stderr)
        return None

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python docx_template_fill_v2.py <template.docx> <content.txt> <output.docx>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]
    
    # Read content
    content_text = read_content_file(content_path)
    
    # Fill template
    docx_path = fill_template_smart(template_path, content_text, output_path)
    
    # Convert to PDF
    pdf_path = convert_to_pdf(docx_path)
    
    print("\n" + "="*60)
    print("GENERATION COMPLETE")
    print("="*60)
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path}")
    print("="*60)
