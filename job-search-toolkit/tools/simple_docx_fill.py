#!/usr/bin/env python3
"""
Simple DOCX Copy-Paste - Just copy content maintaining minimal formatting
"""

import sys
from pathlib import Path
from docx import Document

def main(template_path, content_path, output_path):
    """Simply open template, keep first page header/layout, replace rest"""
    
    print(f"Loading template: {Path(template_path).name}")
    doc = Document(template_path)
    
    # Read content
    with open(content_path, 'r', encoding='utf-8') as f:
        lines = [line.rstrip() for line in f if line.strip()]
    
    print(f"Content lines: {len(lines)}")
    
    # Keep only first 5 paragraphs (header area), delete rest
    while len(doc.paragraphs) > 5:
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    print(f"Kept {len(doc.paragraphs)} header paragraphs, adding content...")
    
    # Add each line as new paragraph
    for line in lines:
        p = doc.add_paragraph()
        
        # Add text with basic formatting
        if line.startswith('•'):
            # Bullet - add with bullet formatting
            run = p.add_run(line)
            p.paragraph_format.left_indent = Pt(20)
        elif line.isupper() and len(line) > 10:
            # Section header - bold and larger
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        else:
            # Normal text
            run = p.add_run(line)
            run.font.size = Pt(11)
    
    # Save
    doc.save(output_path)
    print(f"✓ Saved: {Path(output_path).name}")
    
    # Convert to PDF
    import subprocess
    pdf_path = str(output_path).replace('.docx', '.pdf')
    print("Converting to PDF...")
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', str(Path(output_path).parent),
        str(output_path)
    ], timeout=30, capture_output=True)
    
    if Path(pdf_path).exists():
        print(f"✓ PDF created: {Path(pdf_path).name}")
        # Check pages
        result = subprocess.run(['pdfinfo', pdf_path], capture_output=True, text=True)
        for line in result.stdout.split('\n'):
            if 'Pages:' in line:
                print(f"  {line}")
    
    return output_path

if __name__ == "__main__":
    from docx.shared import Pt
    main(sys.argv[1], sys.argv[2], sys.argv[3])
