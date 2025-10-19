"""
Simple Resume Converter - Markdown to PDF/DOCX
Alternative to pandoc if it's not installed
"""

import os
import sys
from pathlib import Path

def check_dependencies():
    """Check if required packages are installed"""
    missing = []
    
    try:
        import markdown
    except ImportError:
        missing.append("markdown")
    
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate
    except ImportError:
        missing.append("reportlab")
    
    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")
    
    if missing:
        print("❌ Missing dependencies:")
        for pkg in missing:
            print(f"   - {pkg}")
        print("\nInstall with:")
        print(f"   pip install {' '.join(missing)}")
        sys.exit(1)

def convert_to_html(md_file):
    """Convert markdown to HTML"""
    import markdown
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    html = markdown.markdown(md_content, extensions=['extra', 'nl2br'])
    return html

def html_to_pdf(html_content, output_file):
    """Convert HTML to PDF using reportlab"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from bs4 import BeautifulSoup
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create PDF
    # Ensure output_file is a string path for ReportLab compatibility
    output_path = str(output_file)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#1a1a1a',
        spaceAfter=6,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='#2c3e50',
        spaceAfter=6,
        spaceBefore=12
    )
    
    # Process content
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li']):
        text = element.get_text().strip()
        if not text:
            continue
        
        if element.name == 'h1':
            story.append(Paragraph(text, title_style))
        elif element.name in ['h2', 'h3']:
            story.append(Paragraph(text, heading_style))
        elif element.name == 'p':
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        elif element.name == 'li':
            story.append(Paragraph(f"• {text}", styles['Normal']))
    
    doc.build(story)

def markdown_to_docx(md_file, output_file):
    """Convert markdown to DOCX"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(54)  # 0.75 inch
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # H1 - Name
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(26, 26, 26)
        
        # H2 - Sections
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(44, 62, 80)
        
        # H3 - Subsections
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            run = p.runs[0]
            run.font.size = Pt(12)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            run = p.runs[0]
            run.font.size = Pt(11)
        
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        
        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line.replace('**', '').replace('*', '')
            if text and not text.startswith('---'):
                p = doc.add_paragraph(text)
                run = p.runs[0] if p.runs else None
                if run:
                    run.font.size = Pt(11)
    
    # Ensure output_file is a string path for python-docx compatibility
    doc.save(str(output_file))

def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_resume.py <markdown_file>")
        print("\nExample:")
        print("  python convert_resume.py resumes/manpower_19096/Simon_Renauld_Resume_Director_AI_Analytics.md")
        sys.exit(1)
    
    # Check dependencies
    check_dependencies()
    
    # Get input file
    md_file = Path(sys.argv[1])
    if not md_file.exists():
        print(f"❌ File not found: {md_file}")
        sys.exit(1)
    
    # Generate output filenames
    base_name = md_file.stem
    output_dir = md_file.parent
    
    pdf_file = output_dir / f"{base_name}.pdf"
    docx_file = output_dir / f"{base_name}.docx"
    
    print(f"📄 Converting: {md_file.name}")
    print()
    
    # Convert to DOCX
    print("📝 Generating DOCX...")
    try:
        markdown_to_docx(md_file, docx_file)
        print(f"✓ DOCX created: {docx_file}")
        print(f"  Size: {docx_file.stat().st_size / 1024:.1f} KB")
    except Exception as e:
        print(f"✗ Failed to create DOCX: {e}")
    
    print()
    
    # Convert to PDF (requires BeautifulSoup)
    print("📄 Generating PDF...")
    try:
        html = convert_to_html(md_file)
        html_to_pdf(html, pdf_file)
        print(f"✓ PDF created: {pdf_file}")
        print(f"  Size: {pdf_file.stat().st_size / 1024:.1f} KB")
    except Exception as e:
        print(f"✗ Failed to create PDF: {e}")
        if 'BeautifulSoup' in str(e) or 'bs4' in str(e):
            print("  Tip: Install beautifulsoup4 for PDF conversion")
            print("       pip install beautifulsoup4")
        else:
            print("  Tip: If this is a path-related error, ensure the output path is writable and try again.")
    
    print()
    print("✓ Conversion complete!")
    print()
    print("Files generated:")
    print(f"  - {pdf_file} (for viewing/submission)")
    print(f"  - {docx_file} (for ATS parsing)")
    print()
    print("Next steps:")
    print("  1. Open and review both files")
    print("  2. Test DOCX with ATS checker (resume.io, jobscan)")
    print("  3. Ensure formatting looks professional")
    print("  4. Attach to job application")

if __name__ == "__main__":
    main()
