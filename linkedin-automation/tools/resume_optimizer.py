"""
Resume Optimizer - ATS Keyword Optimization Tool
Author: Simon Renauld
Created: October 15, 2025

Purpose: Optimize resume for Applicant Tracking Systems (ATS) by:
1. Extracting keywords from job description
2. Matching against current resume
3. Suggesting keyword additions
4. Generating role-specific resume variants
5. Aligning with simondatalab.de portfolio branding

Usage:
    python tools/resume_optimizer.py \\
        --input ../portfolio-deployment-enhanced/assets/resume/simon-renauld-resume.pdf \\
        --role "Lead Data Engineer" \\
        --job-description path/to/job_description.txt \\
        --output outputs/optimized_resumes/
"""

import argparse
import json
import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Set, Tuple

# Third-party imports (install via requirements.txt)
try:
    import pdfplumber
    from docx import Document
    from docx.shared import Pt, RGBColor
    import spacy
    from sklearn.feature_extraction.text import TfidfVectorizer
    import pandas as pd
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Install requirements: pip install pdfplumber python-docx spacy scikit-learn pandas")
    sys.exit(1)


class ResumeOptimizer:
    """ATS Resume Optimization Engine"""
    
    def __init__(self, config_path: str = "config/portfolio_alignment.json"):
        """Initialize with portfolio alignment config"""
        self.config = self._load_config(config_path)
        self.nlp = spacy.load("en_core_web_sm")  # Requires: python -m spacy download en_core_web_sm
        
    def _load_config(self, config_path: str) -> Dict:
        """Load portfolio alignment configuration"""
        with open(config_path, 'r') as f:
            return json.load(f)
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from resume PDF"""
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_keywords_from_job_description(self, job_desc: str) -> Set[str]:
        """Extract keywords from job description using NLP"""
        doc = self.nlp(job_desc.lower())
        
        # Extract noun chunks (technical terms)
        keywords = set()
        for chunk in doc.noun_chunks:
            keywords.add(chunk.text)
        
        # Extract named entities (tools, technologies)
        for ent in doc.ents:
            if ent.label_ in ["PRODUCT", "ORG", "GPE"]:
                keywords.add(ent.text.lower())
        
        # Common tech keywords
        tech_keywords = [
            "python", "sql", "pyspark", "databricks", "airflow", "spark", "kafka",
            "etl", "elt", "data warehouse", "data engineering", "analytics", "power bi",
            "aws", "azure", "kubernetes", "docker", "ci/cd", "git", "agile"
        ]
        
        for keyword in tech_keywords:
            if keyword in job_desc.lower():
                keywords.add(keyword)
        
        return keywords
    
    def match_keywords(self, resume_text: str, job_keywords: Set[str]) -> Tuple[Set[str], Set[str], float]:
        """Match resume keywords against job description"""
        resume_text_lower = resume_text.lower()
        
        matched = set()
        missing = set()
        
        for keyword in job_keywords:
            if keyword in resume_text_lower:
                matched.add(keyword)
            else:
                missing.add(keyword)
        
        match_score = (len(matched) / len(job_keywords)) * 100 if job_keywords else 0
        
        return matched, missing, match_score
    
    def get_ats_keywords_for_role(self, role: str) -> List[str]:
        """Get predefined ATS keywords for target role"""
        role_map = {
            "lead data engineer": "lead_data_engineer",
            "analytics lead": "analytics_lead",
            "data platform engineer": "data_platform_engineer",
            "head of data": "head_of_data"
        }
        
        role_key = role_map.get(role.lower(), "lead_data_engineer")
        return self.config["ats_keywords"].get(role_key, [])
    
    def generate_optimized_resume(
        self,
        original_resume_text: str,
        role: str,
        missing_keywords: Set[str],
        output_format: str = "docx"
    ) -> str:
        """Generate ATS-optimized resume with missing keywords integrated"""
        
        # Create DOCX (ATS-friendly)
        doc = Document()
        
        # Header: Name & Contact
        header = doc.add_paragraph()
        header.add_run("Simon Renauld\n").bold = True
        header.add_run("Lead Data Engineering & Analytics\n")
        header.add_run("Ho Chi Minh City, Vietnam | +84 923 180 061 | sn.renauld@gmail.com\n")
        header.add_run("Portfolio: simonrenauld.io | LinkedIn: linkedin.com/in/simonrenauld")
        
        # Summary (inject missing keywords naturally)
        summary = doc.add_heading("Professional Summary", level=1)
        summary_text = (
            f"Senior Data Engineering and Analytics leader with 10+ years experience in {role.lower()} roles. "
            f"Specializing in {', '.join(list(missing_keywords)[:5])} with proven track record of "
            f"80% operational automation, 30% faster delivery, and $150K cost savings. "
            f"Expert in building enterprise ETL/DWH platforms, leading cross-functional teams (50+ professionals), "
            f"and implementing data governance frameworks for healthcare and SaaS organizations."
        )
        doc.add_paragraph(summary_text)
        
        # Experience (extract from original resume)
        doc.add_heading("Professional Experience", level=1)
        
        # Add placeholder - in production, parse original resume structure
        doc.add_paragraph("Lead Data Engineering & Analytics — Jio Health (Jan 2022 - Feb 2024)", style='Heading 2')
        doc.add_paragraph("• Coordinated multidisciplinary team of 50+ professionals; improved delivery timelines by 30%")
        doc.add_paragraph("• Automated core processes in Python, reducing manual workload by 80% and saving ~$100K annually")
        doc.add_paragraph("• Architected enterprise ETL/DWH supporting 3x user growth with 50% data quality improvement")
        
        # Skills (inject ALL missing keywords here)
        doc.add_heading("Technical Skills", level=1)
        skills_by_category = {
            "Data Engineering": ["Python", "SQL", "ETL/ELT", "PySpark", "Airflow", "Spark", "Kafka"],
            "Cloud & Infrastructure": ["AWS", "Azure", "Kubernetes", "Docker", "Proxmox"],
            "Databases": ["PostgreSQL", "MySQL", "Redshift", "MongoDB"],
            "Analytics": ["Power BI", "Tableau", "Pandas", "Matplotlib"],
            "Governance": ["Data Quality", "HIPAA", "GDPR", "Compliance"]
        }
        
        # Inject missing keywords into skills
        for category, skills in skills_by_category.items():
            for keyword in missing_keywords:
                if keyword.lower() in ["airflow", "spark", "kafka", "python", "sql"]:
                    if keyword not in skills:
                        skills.append(keyword)
            
            doc.add_paragraph(f"{category}: {', '.join(skills)}")
        
        # Education
        doc.add_heading("Education", level=1)
        doc.add_paragraph("M.Sc. Geography (Geomarketing & GIS) — Laval University, Canada (2014-2015)")
        doc.add_paragraph("B.Sc. Geography (GIS) — Laval University, Canada (2010-2013)")
        
        return doc
    
    def save_resume(self, doc: Document, output_path: str):
        """Save optimized resume to file"""
        doc.save(output_path)
        print(f"✅ Optimized resume saved: {output_path}")
    
    def generate_report(
        self,
        role: str,
        matched: Set[str],
        missing: Set[str],
        match_score: float,
        output_path: str
    ):
        """Generate optimization report"""
        report = {
            "timestamp": datetime.now().isoformat(),
            "target_role": role,
            "ats_match_score": f"{match_score:.1f}%",
            "keywords_matched": sorted(list(matched)),
            "keywords_missing": sorted(list(missing)),
            "recommendations": [
                "Add missing keywords to Skills section",
                "Integrate keywords naturally into Experience bullets",
                "Update summary with role-specific keywords",
                "Ensure DOCX format for ATS compatibility"
            ]
        }
        
        with open(output_path, 'w') as f:
            json.dump(report, f, indent=2)
        
        print(f"\n📊 ATS Optimization Report")
        print(f"Target Role: {role}")
        print(f"Match Score: {match_score:.1f}%")
        print(f"Keywords Matched: {len(matched)}")
        print(f"Keywords Missing: {len(missing)}")
        print(f"\nMissing Keywords: {', '.join(sorted(list(missing))[:10])}")
        print(f"\n✅ Full report saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="ATS Resume Optimizer")
    parser.add_argument("--input", required=True, help="Path to resume PDF")
    parser.add_argument("--role", required=True, help="Target role (e.g., 'Lead Data Engineer')")
    parser.add_argument("--job-description", help="Path to job description text file")
    parser.add_argument("--output", default="outputs/optimized_resumes/", help="Output directory")
    
    args = parser.parse_args()
    
    # Initialize optimizer
    optimizer = ResumeOptimizer()
    
    # Extract resume text
    print(f"📄 Extracting text from: {args.input}")
    resume_text = optimizer.extract_text_from_pdf(args.input)
    
    # Get keywords
    if args.job_description:
        with open(args.job_description, 'r') as f:
            job_desc = f.read()
        job_keywords = optimizer.extract_keywords_from_job_description(job_desc)
    else:
        # Use predefined keywords for role
        job_keywords = set(optimizer.get_ats_keywords_for_role(args.role))
    
    # Match keywords
    print(f"🔍 Analyzing keywords for role: {args.role}")
    matched, missing, match_score = optimizer.match_keywords(resume_text, job_keywords)
    
    # Generate optimized resume
    print(f"⚙️  Generating optimized resume...")
    optimized_doc = optimizer.generate_optimized_resume(
        resume_text,
        args.role,
        missing,
        output_format="docx"
    )
    
    # Save outputs
    os.makedirs(args.output, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    role_slug = args.role.lower().replace(" ", "_")
    
    resume_path = os.path.join(args.output, f"resume_{role_slug}_{timestamp}.docx")
    report_path = os.path.join(args.output, f"report_{role_slug}_{timestamp}.json")
    
    optimizer.save_resume(optimized_doc, resume_path)
    optimizer.generate_report(args.role, matched, missing, match_score, report_path)
    
    print(f"\n✨ Optimization complete!")
    print(f"Resume: {resume_path}")
    print(f"Report: {report_path}")


if __name__ == "__main__":
    main()
